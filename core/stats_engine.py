"""
stats_engine.py — Modele de donnees et persistance pour les statistiques

Structure sur disque :
  data/projects/{slug}/stats/
    doc_stats/
        index.json      -- [{doc_id, title, path, injections:[{date, word_count}], baseline_wpd}]
    custom_stats/
        index.json      -- [{stat_id, name, chart_type, data:{labels,values,colors?}, description, created_at}]

API publique :
  store = StatsStore(project_dir)
  # Docs
  entry = store.add_doc_injection(path, word_count, baseline_wpd=None)
  store.update_baseline(doc_id, words_per_day)
  store.delete_doc_stat(doc_id)
  entries = store.list_doc_stats()      -> [DocStatEntry, ...]
  entry   = store.get_doc_stat(doc_id)  -> DocStatEntry | None
  # Custom
  entry = store.add_custom_stat(name, chart_type, data, description)
  store.delete_custom_stat(stat_id)
  entries = store.list_custom_stats()   -> [CustomStatEntry, ...]
  # Utilitaire
  wc = count_words_docx(path)           -> int
"""

import hashlib
import json
import logging
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class DocInjection:
    date: str        # ISO-8601 date : "2026-05-21"
    word_count: int


@dataclass
class DocStatEntry:
    doc_id: str                           # sha256 du chemin normalisé (12 premiers hex)
    title: str                            # nom du fichier (basename)
    path: str                             # chemin absolu
    injections: list[DocInjection] = field(default_factory=list)
    baseline_wpd: int | None = None       # mots/jour de base fournis par l'utilisateur

    # ---- sérialization --------------------------------------------------

    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "title": self.title,
            "path": self.path,
            "injections": [{"date": i.date, "word_count": i.word_count} for i in self.injections],
            "baseline_wpd": self.baseline_wpd,
        }

    @classmethod
    def from_dict(cls, d: dict) -> "DocStatEntry":
        injections = [DocInjection(date=i["date"], word_count=i["word_count"]) for i in d.get("injections", [])]
        return cls(
            doc_id=d["doc_id"],
            title=d["title"],
            path=d["path"],
            injections=injections,
            baseline_wpd=d.get("baseline_wpd"),
        )

    @property
    def latest_word_count(self) -> int | None:
        if not self.injections:
            return None
        return self.injections[-1].word_count

    @property
    def word_count_delta(self) -> int | None:
        """Différence entre dernière et avant-dernière injection."""
        if len(self.injections) < 2:
            return None
        return self.injections[-1].word_count - self.injections[-2].word_count


@dataclass
class CustomStatEntry:
    stat_id: str
    name: str
    chart_type: str   # "bar" | "line" | "pie"
    data: dict        # {"labels": [...], "values": [...], "colors": [...] (optionnel)}
    description: str
    created_at: str   # ISO-8601

    def to_dict(self) -> dict:
        return {
            "stat_id": self.stat_id,
            "name": self.name,
            "chart_type": self.chart_type,
            "data": self.data,
            "description": self.description,
            "created_at": self.created_at,
        }

    @classmethod
    def from_dict(cls, d: dict) -> "CustomStatEntry":
        return cls(
            stat_id=d["stat_id"],
            name=d["name"],
            chart_type=d["chart_type"],
            data=d["data"],
            description=d.get("description", ""),
            created_at=d["created_at"],
        )

    @property
    def summary_value(self) -> str:
        """Valeur résumée pour l'affichage dans la liste (ex: '23 %' ou '3 éléments')."""
        values = self.data.get("values", [])
        if not values:
            return "—"
        if self.chart_type == "pie":
            # On affiche la valeur max avec son label
            max_idx = max(range(len(values)), key=lambda i: values[i])
            labels = self.data.get("labels", [])
            label = labels[max_idx] if max_idx < len(labels) else "?"
            return f"{label}: {values[max_idx]:.0f} %"
        total = sum(values)
        return f"{total:,.0f}" if total >= 1000 else f"{total}"


# ---------------------------------------------------------------------------
# StatsStore
# ---------------------------------------------------------------------------

class StatsStore:
    """Gestion de la persistance des statistiques d'un projet."""

    def __init__(self, project_dir: Path | str):
        self._project_dir = Path(project_dir)
        self._doc_index_path = self._project_dir / "stats" / "doc_stats" / "index.json"
        self._custom_index_path = self._project_dir / "stats" / "custom_stats" / "index.json"
        self._doc_index_path.parent.mkdir(parents=True, exist_ok=True)
        self._custom_index_path.parent.mkdir(parents=True, exist_ok=True)
        self._doc_stats: dict[str, DocStatEntry] = {}
        self._custom_stats: dict[str, CustomStatEntry] = {}
        self._load()

    # ------------------------------------------------------------------
    # Chargement / sauvegarde
    # ------------------------------------------------------------------

    def _load(self) -> None:
        if self._doc_index_path.exists():
            raw = json.loads(self._doc_index_path.read_text(encoding="utf-8"))
            for d in raw:
                entry = DocStatEntry.from_dict(d)
                self._doc_stats[entry.doc_id] = entry
        if self._custom_index_path.exists():
            raw = json.loads(self._custom_index_path.read_text(encoding="utf-8"))
            for d in raw:
                entry = CustomStatEntry.from_dict(d)
                self._custom_stats[entry.stat_id] = entry

    def _save_doc_index(self) -> None:
        data = [e.to_dict() for e in self._doc_stats.values()]
        self._doc_index_path.write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    def _save_custom_index(self) -> None:
        data = [e.to_dict() for e in self._custom_stats.values()]
        self._custom_index_path.write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    # ------------------------------------------------------------------
    # Utilitaire — ID stable pour un chemin
    # ------------------------------------------------------------------

    @staticmethod
    def doc_id_for_path(path: str | Path) -> str:
        normalized = str(Path(path).resolve()).lower()
        return hashlib.sha256(normalized.encode()).hexdigest()[:16]

    # ------------------------------------------------------------------
    # API — Documents
    # ------------------------------------------------------------------

    def get_doc_stat(self, doc_id: str) -> DocStatEntry | None:
        return self._doc_stats.get(doc_id)

    def get_doc_stat_by_path(self, path: str | Path) -> DocStatEntry | None:
        doc_id = self.doc_id_for_path(path)
        return self._doc_stats.get(doc_id)

    def add_doc_injection(
        self,
        path: str | Path,
        word_count: int,
        baseline_wpd: int | None = None,
    ) -> DocStatEntry:
        """Enregistre une injection (nouvelle ou mise à jour).
        Si le doc est connu, ajoute l'injection. Sinon le crée.
        baseline_wpd n'est appliqué que si fourni et pas encore défini.
        """
        doc_id = self.doc_id_for_path(path)
        today = datetime.now().strftime("%Y-%m-%d")

        if doc_id in self._doc_stats:
            entry = self._doc_stats[doc_id]
            # Mise à jour du chemin si le fichier a été déplacé
            entry.path = str(path)
            # Ajouter l'injection seulement si elle n'existe pas déjà pour today
            dates_existing = {inj.date for inj in entry.injections}
            if today not in dates_existing:
                entry.injections.append(DocInjection(date=today, word_count=word_count))
            else:
                # Mettre à jour le compte du jour si réinjecté dans la même journée
                for inj in entry.injections:
                    if inj.date == today:
                        inj.word_count = word_count
                        break
            if baseline_wpd is not None and entry.baseline_wpd is None:
                entry.baseline_wpd = baseline_wpd
        else:
            entry = DocStatEntry(
                doc_id=doc_id,
                title=Path(path).name,
                path=str(path),
                injections=[DocInjection(date=today, word_count=word_count)],
                baseline_wpd=baseline_wpd,
            )
            self._doc_stats[doc_id] = entry

        self._save_doc_index()
        return entry

    def update_baseline(self, doc_id: str, words_per_day: int) -> None:
        if doc_id not in self._doc_stats:
            raise KeyError(f"StatsStore.update_baseline : doc_id inconnu '{doc_id}'")
        self._doc_stats[doc_id].baseline_wpd = words_per_day
        self._save_doc_index()

    def delete_doc_stat(self, doc_id: str) -> None:
        if doc_id not in self._doc_stats:
            raise KeyError(f"StatsStore.delete_doc_stat : doc_id inconnu '{doc_id}'")
        del self._doc_stats[doc_id]
        self._save_doc_index()

    def list_doc_stats(self) -> list[DocStatEntry]:
        return list(self._doc_stats.values())

    # ------------------------------------------------------------------
    # API — Stats personnalisées
    # ------------------------------------------------------------------

    def add_custom_stat(
        self,
        name: str,
        chart_type: str,
        data: dict,
        description: str = "",
    ) -> CustomStatEntry:
        stat_id = uuid.uuid4().hex[:12]
        entry = CustomStatEntry(
            stat_id=stat_id,
            name=name,
            chart_type=chart_type,
            data=data,
            description=description,
            created_at=datetime.now().isoformat(timespec="seconds"),
        )
        self._custom_stats[stat_id] = entry
        self._save_custom_index()
        return entry

    def delete_custom_stat(self, stat_id: str) -> None:
        if stat_id not in self._custom_stats:
            raise KeyError(f"StatsStore.delete_custom_stat : stat_id inconnu '{stat_id}'")
        del self._custom_stats[stat_id]
        self._save_custom_index()

    def list_custom_stats(self) -> list[CustomStatEntry]:
        return list(self._custom_stats.values())


# ---------------------------------------------------------------------------
# Utilitaire — comptage mots dans un .docx
# ---------------------------------------------------------------------------

def count_words_docx(path: str | Path) -> int:
    """Compte le nombre de mots dans un .docx (tous paragraphes, toutes tables)."""
    from docx import Document  # noqa: PLC0415
    doc = Document(str(path))
    total = 0
    for para in doc.paragraphs:
        total += len(para.text.split())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += len(para.text.split())
    return total
