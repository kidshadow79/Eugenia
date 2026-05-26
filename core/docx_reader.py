"""
docx_reader.py — Extraction de texte depuis un document (multi-format)

Formats supportés :
    .docx  — python-docx
    .txt   — UTF-8 (fallback latin-1)
    .md    — identique .txt
    .pdf   — pypdf (pip install pypdf)
    .json  — dump indenté du contenu sérialisable

Responsabilité unique : retourner du texte brut propre prêt pour le ChunkManager.

Usage :
    from core.docx_reader import read_document, estimate_word_count
    text = read_document(Path("mon_roman.docx"))
"""

import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Extensions acceptées
SUPPORTED_EXTENSIONS = {".docx", ".txt", ".md", ".pdf", ".json"}


# ─── Lecteurs par format ──────────────────────────────────────────────────────

def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx est requis : pip install python-docx")
    try:
        doc = Document(str(path))
    except Exception as e:
        raise ValueError(f"Impossible d'ouvrir {path.name} : {e}") from e

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        raise ValueError(f"{path.name} ne contient aucun texte extractible.")
    return "\n\n".join(paragraphs)


def _read_txt(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = path.read_text(encoding=enc)
            text = text.strip()
            if not text:
                raise ValueError(f"{path.name} est vide.")
            return text
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Impossible de décoder {path.name} (UTF-8 / latin-1 échoués).")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf est requis : pip install pypdf")
    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise ValueError(f"Impossible d'ouvrir {path.name} : {e}") from e

    pages: list[str] = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            pages.append(t)
    if not pages:
        raise ValueError(f"{path.name} ne contient aucun texte extractible (PDF scanné ?).")
    return "\n\n".join(pages)


def _read_json(path: Path) -> str:
    try:
        raw = path.read_text(encoding="utf-8")
        data = json.loads(raw)
    except (json.JSONDecodeError, OSError) as e:
        raise ValueError(f"JSON invalide dans {path.name} : {e}") from e

    # Convertir en texte lisible : dump indenté
    text = json.dumps(data, ensure_ascii=False, indent=2)
    if not text.strip():
        raise ValueError(f"{path.name} est vide.")
    return text


# ─── API publique ─────────────────────────────────────────────────────────────

def read_document(path: Path) -> str:
    """
    Extrait le texte brut d'un document (docx / txt / md / pdf / json).

    Returns:
        Texte nettoyé prêt pour le ChunkManager.
        Lève ValueError si le fichier ne peut pas être lu.
        Lève ImportError si une dépendance optionnelle est manquante.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = _read_docx(path)
    elif ext in (".txt", ".md"):
        text = _read_txt(path)
    elif ext == ".pdf":
        text = _read_pdf(path)
    elif ext == ".json":
        text = _read_json(path)
    else:
        raise ValueError(
            f"Format non supporté : {ext}. "
            f"Formats acceptés : {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    logger.info("docx_reader — %s : %d chars extraits", path.name, len(text))
    return text


# Alias de rétrocompatibilité — l'ancien nom est encore utilisé dans ingest_worker
def read_docx(path: Path) -> str:
    """Alias rétrocompat → appelle read_document."""
    return read_document(path)


def estimate_word_count(text: str) -> int:
    """Estimation rapide du nombre de mots (utile pour l'affichage avant import)."""
    return len(text.split())

